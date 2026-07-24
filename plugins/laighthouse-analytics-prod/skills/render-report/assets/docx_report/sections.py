"""python-docx section renderers consumed by build.py."""
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

import charts

_HEADER_FILL = "F1F5F9"
_TOTAL_ROW_FILL = "F1F5F9"
_BAR_TRACK_FILL = "E2E8F0"
_BAR_WIDTH_EMU = Emu(914400)  # 1 inch, matches the fixed-width CSS bar track
_TOTAL_ROW_LABELS = {"합계", "총계"}


def add_title(document, title, period=None):
    document.add_heading(title, level=1)
    if period:
        p = document.add_paragraph(f"기간: {period}")
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def add_heading(document, text):
    document.add_heading(text, level=2)


def _diff_color(diff_value):
    if diff_value is None:
        return None
    if diff_value > 0:
        return RGBColor(0x16, 0xA3, 0x4A)
    if diff_value < 0:
        return RGBColor(0xDC, 0x26, 0x26)
    return RGBColor(0x6B, 0x72, 0x80)


def _shade_cell(cell, hex_color):
    """Set a table cell's background fill (CSS `background` equivalent)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    cell._tc.get_or_add_tcPr().append(shd)


def _is_bar_cell(value):
    return isinstance(value, dict) and value.get("type") == "bar"


def _render_bar_cell(cell, pct, color, label):
    """Approximate a CSS `width:{pct}%` progress bar with a 2-cell nested
    table (filled + track), since Word has no native proportional-width bar.
    """
    pct = max(0.0, min(float(pct), 100.0))
    cell.text = ""
    bar_table = cell.add_table(rows=1, cols=2)
    bar_table.autofit = False
    filled_width = Emu(int(_BAR_WIDTH_EMU * pct / 100))
    track_width = _BAR_WIDTH_EMU - filled_width
    filled_cell, track_cell = bar_table.rows[0].cells
    bar_table.columns[0].width = filled_width
    bar_table.columns[1].width = track_width
    filled_cell.width = filled_width
    track_cell.width = track_width
    _shade_cell(filled_cell, color)
    _shade_cell(track_cell, _BAR_TRACK_FILL)
    cell.add_paragraph(label)


def add_kpi_cards(document, cards):
    table = document.add_table(rows=2, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    label_cells = table.rows[0].cells
    value_cells = table.rows[1].cells
    for col, card in enumerate(cards):
        label_cells[col].text = card["label"]

        value_cell = value_cells[col]
        paragraph = value_cell.paragraphs[0]

        value_run = paragraph.add_run(card["value"])
        value_run.font.size = Pt(14)
        value_run.font.bold = True
        accent = card.get("accent")
        if accent:
            value_run.font.color.rgb = RGBColor.from_string(accent.lstrip("#").upper())

        if card.get("diff"):
            diff_run = paragraph.add_run(f"  ({card['diff']})")
            diff_run.font.size = Pt(14)
            diff_run.font.bold = True
            color = _diff_color(card.get("diff_value"))
            if color is not None:
                diff_run.font.color.rgb = color


def add_data_table(document, heading, headers, rows):
    if heading:
        add_heading(document, heading)
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # table.cell(r, c) rescans every cell in the table on each call; for
    # large tables (e.g. 1000+ keyword rows) that's O(rows*cols) per cell,
    # O(n^2) overall. table.rows[i].cells only scans that row.
    header_cells = table.rows[0].cells
    for col, header in enumerate(headers):
        header_cells[col].text = header
        header_cells[col].paragraphs[0].runs[0].font.bold = True
        _shade_cell(header_cells[col], _HEADER_FILL)
    for row_idx, row in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        is_total_row = bool(row) and str(row[0]) in _TOTAL_ROW_LABELS
        for col, value in enumerate(row):
            cell = row_cells[col]
            if _is_bar_cell(value):
                _render_bar_cell(cell, value["pct"], value.get("color", "3B82F6"), value.get("label", f"{value['pct']}%"))
            else:
                cell.text = str(value)
            if is_total_row:
                _shade_cell(cell, _TOTAL_ROW_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True


def add_text_section(document, heading, body):
    if heading:
        add_heading(document, heading)
    document.add_paragraph(body)


def add_combo_chart_section(document, heading, categories, bar_series, line_series):
    if heading:
        add_heading(document, heading)
    charts.add_combo_chart(document, categories, bar_series, line_series)
