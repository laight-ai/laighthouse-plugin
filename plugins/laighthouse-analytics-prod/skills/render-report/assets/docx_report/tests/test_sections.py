import time

from docx import Document
from docx.oxml.ns import qn

import sections as sec


def _cell_fill(cell):
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def test_add_kpi_cards_renders_label_and_value(tmp_path):
    document = Document()
    sec.add_kpi_cards(document, [
        {"label": "월 예산대비 소진율", "value": "62.3%", "diff": "-7.1%p", "diff_value": -7.1},
        {"label": "월 목표 매출 대비 달성률", "value": "96.6%", "diff": "+3.3%p", "diff_value": 3.3},
    ])

    out_path = tmp_path / "kpi.docx"
    document.save(out_path)

    reopened = Document(out_path)
    table = reopened.tables[0]
    assert table.cell(0, 0).text == "월 예산대비 소진율"
    assert "62.3%" in table.cell(1, 0).text
    assert "-7.1%p" in table.cell(1, 0).text


def test_add_kpi_cards_applies_accent_color_to_value(tmp_path):
    document = Document()
    sec.add_kpi_cards(document, [
        {"label": "월 예산 목표", "value": "15,000,000", "accent": "#3b82f6"},
    ])

    table = document.tables[0]
    value_run = table.cell(1, 0).paragraphs[0].runs[0]
    assert value_run.font.color.rgb.__str__().upper() == "3B82F6"


def test_add_data_table_shades_header_row(tmp_path):
    document = Document()
    sec.add_data_table(
        document,
        heading=None,
        headers=["캠페인", "매출"],
        rows=[["a", "1"]],
    )

    table = document.tables[0]
    assert _cell_fill(table.cell(0, 0)) == "F1F5F9"
    assert _cell_fill(table.cell(1, 0)) is None


def test_add_data_table_bolds_and_shades_total_row(tmp_path):
    document = Document()
    sec.add_data_table(
        document,
        heading=None,
        headers=["매체", "소진율"],
        rows=[["네이버 파워링크", "50.4%"], ["합계", "55.9%"]],
    )

    table = document.tables[0]
    total_cell = table.cell(2, 0)
    assert _cell_fill(total_cell) == "F1F5F9"
    assert total_cell.paragraphs[0].runs[0].font.bold is True
    # non-total data row stays unshaded/unbolded
    assert _cell_fill(table.cell(1, 0)) is None


def test_add_data_table_renders_progress_bar_cell(tmp_path):
    document = Document()
    sec.add_data_table(
        document,
        heading=None,
        headers=["매체", "예산 소진율"],
        rows=[["네이버 브랜드검색", {"type": "bar", "pct": 55.9, "color": "3B82F6", "label": "55.9%"}]],
    )

    out_path = tmp_path / "bar.docx"
    document.save(out_path)
    reopened = Document(out_path)

    bar_cell = reopened.tables[0].cell(1, 1)
    assert "55.9%" in bar_cell.text
    nested_table = bar_cell.tables[0]
    filled_cell, track_cell = nested_table.rows[0].cells
    assert _cell_fill(filled_cell) == "3B82F6"
    assert _cell_fill(track_cell) == sec._BAR_TRACK_FILL


def test_add_data_table_renders_headers_and_rows(tmp_path):
    document = Document()
    sec.add_data_table(
        document,
        heading="캠페인별 성과",
        headers=["캠페인", "매출", "ROAS"],
        rows=[["05_GT케이(SPBR)_MO", "1,320,543", "1017%"]],
    )

    out_path = tmp_path / "table.docx"
    document.save(out_path)

    reopened = Document(out_path)
    table = reopened.tables[0]
    assert table.cell(0, 0).text == "캠페인"
    assert table.cell(1, 0).text == "05_GT케이(SPBR)_MO"
    assert table.cell(1, 2).text == "1017%"


def test_add_data_table_stays_fast_for_large_tables(tmp_path):
    # Regression guard: table.cell(r, c) rescans the whole table on every
    # call, making it O(n^2) for large tables (e.g. keyword performance with
    # 1000+ rows took effectively forever). row.cells access must stay row-
    # scoped for this to remain linear.
    headers = ["키워드", "노출", "클릭", "광고비", "CPC", "클릭율", "CPM", "구매건수", "매출", "ROAS"]
    rows = [[f"keyword-{i}", i, i, i, i, i, i, i, i, i] for i in range(1500)]

    document = Document()
    started = time.monotonic()
    sec.add_data_table(document, heading="키워드별 성과", headers=headers, rows=rows)
    elapsed = time.monotonic() - started

    assert elapsed < 5, f"add_data_table took {elapsed:.1f}s for 1500 rows, expected a few seconds"
    table = document.tables[0]
    assert table.rows[1].cells[0].text == "keyword-0"
    assert table.rows[1500].cells[0].text == "keyword-1499"


def test_add_text_section_renders_heading_and_body(tmp_path):
    document = Document()
    sec.add_text_section(document, "Executive Summary", "ROAS가 목표 대비 118% 달성되었습니다.")

    out_path = tmp_path / "text.docx"
    document.save(out_path)

    reopened = Document(out_path)
    full_text = "\n".join(p.text for p in reopened.paragraphs)
    assert "Executive Summary" in full_text
    assert "ROAS가 목표 대비 118% 달성되었습니다." in full_text
