import json
import sys

from docx import Document

import build


SAMPLE_DATA = {
    "title": "다형식품 MTD 보고서",
    "period": "2026-05-01 ~ 2026-05-15",
    "sections": [
        {"type": "heading", "text": "목표 달성 현황"},
        {
            "type": "kpi_cards",
            "cards": [
                {"label": "월 예산대비 소진율", "value": "62.3%", "diff": "-7.1%p", "diff_value": -7.1},
                {"label": "월 누적 ROAS", "value": "506%", "diff": "+42%p", "diff_value": 42},
            ],
        },
        {
            "type": "chart",
            "heading": "월별 광고 성과",
            "categories": ["25년 11월", "25년 12월", "26년 1월"],
            "bar_series": [
                {"name": "광고비", "values": [1000000, 1100000, 1200000]},
                {"name": "매출", "values": [5000000, 5200000, 5400000]},
            ],
            "line_series": {"name": "ROAS", "values": [500, 473, 450]},
        },
        {
            "type": "table",
            "heading": "캠페인별 성과",
            "headers": ["캠페인", "매출", "ROAS"],
            "rows": [["05_GT케이(SPBR)_MO", "1,320,543", "1017%"]],
        },
        {
            "type": "text",
            "heading": "Executive Summary",
            "body": "ROAS가 목표 대비 118% 달성되었습니다.",
        },
    ],
}


def test_build_document_creates_all_section_types():
    document = build.build_document(SAMPLE_DATA)
    assert len(document.tables) == 2  # kpi_cards + table
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "다형식품 MTD 보고서" in full_text
    assert "Executive Summary" in full_text


def test_build_document_rejects_unknown_section_type():
    bad_data = {"title": "t", "sections": [{"type": "nope"}]}
    try:
        build.build_document(bad_data)
        assert False, "expected ValueError"
    except ValueError as e:
        assert "nope" in str(e)


def test_main_writes_docx_file(tmp_path):
    data_path = tmp_path / "data.json"
    data_path.write_text(json.dumps(SAMPLE_DATA, ensure_ascii=False), encoding="utf-8")
    out_path = tmp_path / "out.docx"

    old_argv = sys.argv
    sys.argv = ["build.py", "--data", str(data_path), "--out", str(out_path)]
    try:
        build.main()
    finally:
        sys.argv = old_argv

    assert out_path.exists()
    reopened = Document(out_path)
    assert len(reopened.tables) == 2
