import zipfile

from docx import Document

import charts


def test_add_combo_chart_creates_chart_part_and_reference(tmp_path):
    document = Document()
    document.add_paragraph("placeholder")

    rId = charts.add_combo_chart(
        document,
        categories=["1월", "2월", "3월"],
        bar_series=[
            {"name": "광고비", "values": [100, 200, 150]},
            {"name": "매출", "values": [300, 400, 350]},
        ],
        line_series={"name": "ROAS", "values": [300, 200, 233]},
        title="월별 광고 성과",
    )

    out_path = tmp_path / "chart_test.docx"
    document.save(out_path)

    with zipfile.ZipFile(out_path) as z:
        names = z.namelist()
        assert "word/charts/chart1.xml" in names
        chart_xml = z.read("word/charts/chart1.xml").decode("utf-8")
        document_xml = z.read("word/document.xml").decode("utf-8")

    assert "광고비" in chart_xml
    assert "ROAS" in chart_xml
    assert "1월" in chart_xml
    assert "barChart" in chart_xml
    assert "lineChart" in chart_xml
    assert rId in document_xml
    assert "drawingml/2006/chart" in document_xml


def test_add_combo_chart_supports_multiple_charts_in_one_document(tmp_path):
    document = Document()

    charts.add_combo_chart(
        document,
        categories=["1월", "2월"],
        bar_series=[{"name": "광고비", "values": [10, 20]}],
        line_series={"name": "ROAS", "values": [100, 110]},
    )
    charts.add_combo_chart(
        document,
        categories=["3월", "4월"],
        bar_series=[{"name": "광고비", "values": [30, 40]}],
        line_series={"name": "ROAS", "values": [120, 130]},
    )

    out_path = tmp_path / "two_charts.docx"
    document.save(out_path)

    with zipfile.ZipFile(out_path) as z:
        names = z.namelist()
        assert "word/charts/chart1.xml" in names
        assert "word/charts/chart2.xml" in names
