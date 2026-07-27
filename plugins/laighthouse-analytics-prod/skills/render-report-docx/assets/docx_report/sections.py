"""python-docx section renderers consumed by build.py.

Written fresh for this skill (no reuse of the reverted first-generation
renderer): every component styles itself from theme.py — banner section
headers with an accent bar, tinted KPI cards, borderless header-band tables
with content-proportional column widths, and a footer with live page
numbers. Display rounding happens once in _prettify(), and large tables
drop 매출(gross) 0원 rows before rendering.
"""
import base64
import io
import re

from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

import charts
import theme

_TOTAL_ROW_LABELS = {"합계", "총계"}
_GROSS_HEADERS = {"매출", "전환매출", "매출액"}
_BAR_WIDTH_EMU = Emu(914400)  # 1 inch, matches the fixed-width CSS bar track


# ── text helpers ───────────────────────────────────────────────────────

def _rgb(hex_color):
    return RGBColor.from_string(hex_color.lstrip("#").upper())


_LONG_DECIMAL = re.compile(r"(\d+)\.(\d{3,})")


def _prettify(text):
    """Round runaway float precision for display: 55.88918788518661 → 55.89.
    Only decimal runs of 3+ places are touched, so pre-formatted values
    (2,449.08 / 727% / dates) pass through unchanged."""
    def _round(match):
        rounded = round(float(f"{match.group(1)}.{match.group(2)}"), 2)
        return f"{rounded:.2f}".rstrip("0").rstrip(".")
    return _LONG_DECIMAL.sub(_round, str(text))


def _style_run(run, size=None, color=None, bold=None):
    run.font.name = theme.FONT
    # Korean glyphs resolve through the eastAsia slot, not the latin one
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), theme.FONT)
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = _rgb(color)
    if bold is not None:
        run.font.bold = bold
    return run


def _para(container, text="", size=None, color=None, bold=None,
          space_before=Pt(0), space_after=Pt(0), align=None, first=False):
    if first and container.paragraphs and not container.paragraphs[0].runs:
        p = container.paragraphs[0]
    else:
        p = container.add_paragraph()
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if align is not None:
        p.alignment = align
    if text != "":
        _style_run(p.add_run(), size=size, color=color, bold=bold).text = _prettify(text)
    return p


# ── oxml helpers ───────────────────────────────────────────────────────

def _edge(name, val="single", sz=4, color="auto"):
    el = OxmlElement(f"w:{name}")
    el.set(qn("w:val"), val)
    if val != "nil":
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
    return el


def _shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    cell._tc.get_or_add_tcPr().append(shd)


def _set_table_borders(table, **edges):
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        spec = edges.get(name, "nil")
        if spec == "nil":
            borders.append(_edge(name, "nil"))
        else:
            val, sz, color = spec
            borders.append(_edge(name, val, sz, color))
    table._tbl.tblPr.append(borders)


def _set_cell_border(cell, name, val="single", sz=4, color="auto"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    borders.append(_edge(name, val, sz, color))


def _set_cell_margins(table, top=80, bottom=80, left=110, right=110):
    mar = OxmlElement("w:tblCellMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    table._tbl.tblPr.append(mar)


def _fixed_layout(table, col_widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    for idx, width in enumerate(col_widths):
        table.columns[idx].width = width
        for cell in table.columns[idx].cells:
            cell.width = width


# ── document scaffold ──────────────────────────────────────────────────

def setup_document(document):
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = theme.PAGE_W
    section.page_height = theme.PAGE_H
    section.top_margin = section.bottom_margin = theme.MARGIN
    section.left_margin = section.right_margin = theme.MARGIN

    normal = document.styles["Normal"]
    normal.font.name = theme.FONT
    normal.font.size = theme.SIZE_BODY
    normal.font.color.rgb = _rgb(theme.TEXT)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), theme.FONT)
    normal.paragraph_format.space_after = Pt(0)


def add_title(document, title, period=None):
    _para(document, title, size=theme.SIZE_TITLE, color=theme.TEXT_STRONG,
          bold=True, first=True)
    p = _para(document, f"📅 {period}" if period else "",
              size=theme.SIZE_PERIOD, color=theme.TEXT_MUTED,
              space_before=Pt(2), space_after=Pt(8))
    # thick accent rule closing the header block
    pbdr = OxmlElement("w:pBdr")
    pbdr.append(_edge("bottom", "single", 18, theme.ACCENT))
    p.paragraph_format.element.get_or_add_pPr().append(pbdr)


def add_heading(document, text, number=None, page_break=False):
    """Section header: full-width tinted banner with an accent left bar and
    an accent-colored section number — the .section-title translated into a
    print-document device. page_break starts the section on a fresh page."""
    p = _para(document, "", space_before=Pt(18), space_after=Pt(10))
    if page_break:
        p.paragraph_format.page_break_before = True
    if number:
        _style_run(p.add_run(), size=theme.SIZE_SECTION_NO,
                   color=theme.ACCENT, bold=True).text = f"{number}"
        _style_run(p.add_run(), size=theme.SIZE_SECTION, bold=True).text = "  "
    _style_run(p.add_run(), size=theme.SIZE_SECTION, color=theme.TEXT_STRONG,
               bold=True).text = _prettify(text)
    ppr = p.paragraph_format.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), theme.FILL_BANNER)
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    pbdr.append(_edge("left", "single", 20, theme.ACCENT))
    ppr.append(pbdr)
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p


def add_footer(document):
    """Centered brand line + live page number at the right edge."""
    footer = document.sections[0].footer
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Emu(int(theme.CONTENT_W) // 2),
                                              WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Emu(int(theme.CONTENT_W)),
                                              WD_TAB_ALIGNMENT.RIGHT)
    _style_run(p.add_run("\t"), size=theme.SIZE_FOOTER)
    _style_run(p.add_run("Engineered by Laighthouse AI"),
               size=theme.SIZE_FOOTER, color=theme.TEXT_FAINT)
    _style_run(p.add_run("\t"), size=theme.SIZE_FOOTER)
    fld = parse_xml(
        '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:instr=" PAGE "><w:r><w:rPr><w:color w:val="94A3B8"/>'
        '<w:sz w:val="17"/></w:rPr><w:t>1</w:t></w:r></w:fldSimple>')
    p._p.append(fld)


# ── KPI cards ──────────────────────────────────────────────────────────

_PCT_VALUE = re.compile(r"^(\d+(?:\.\d+)?)%$")
_GOAL_IN_DIFF = re.compile(r"목표\s+([\d,]+(?:\.\d+)?)%")


def _progress_pct(card):
    """(bar_fill_pct, achieved_pct) for a percentage KPI card, or None.
    Rate cards (소진율/달성률) use the value itself; the 누적 ROAS card
    compares value vs the 목표 in its diff line. The fill caps at 100% but
    the achieved label keeps the real ratio (131% 초과 달성 등). Static
    goal cards (diff 없음 — 예: 월 ROAS 목표) get no bar: there is no
    progress to show."""
    if not card.get("diff"):
        return None
    match = _PCT_VALUE.match(str(card.get("value", "")).strip())
    if not match:
        return None
    value = float(match.group(1))
    goal_match = _GOAL_IN_DIFF.search(str(card.get("diff", "")))
    if goal_match:
        goal = float(goal_match.group(1).replace(",", ""))
        if goal > 0:
            achieved = value / goal * 100
            return min(achieved, 100.0), achieved
    return min(value, 100.0), value


def _card_progress_bar(cell, pct, color, card_w):
    """Slim progress bar under the KPI value — 목표 대비 진척을 한눈에."""
    bar_w = int(card_w * 0.72)
    bar = cell.add_table(rows=1, cols=2)
    bar.autofit = False
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    filled = Emu(max(int(bar_w * pct / 100), 1))
    track = Emu(bar_w - int(filled))
    filled_cell, track_cell = bar.rows[0].cells
    bar.columns[0].width = filled
    bar.columns[1].width = track
    filled_cell.width, track_cell.width = filled, track
    _shade_cell(filled_cell, color)
    _shade_cell(track_cell, theme.BORDER)


def add_kpi_cards(document, cards):
    """Tinted, bordered card cells separated by spacer columns — the HTML
    flex card row. Percentage cards (소진율/달성률/누적 ROAS) get a slim
    progress bar under the value so 목표 대비 진척이 시각적으로 읽힌다."""
    n = len(cards)
    cols = 2 * n - 1
    gap = Cm(0.3)
    card_w = Emu(int((theme.CONTENT_W - gap * (n - 1)) / n))
    table = document.add_table(rows=1, cols=cols)
    _set_table_borders(table)  # all nil; card edges are per-cell
    _set_cell_margins(table, top=150, bottom=150, left=110, right=110)
    _fixed_layout(table, [card_w if i % 2 == 0 else gap for i in range(cols)])

    for i, card in enumerate(cards):
        cell = table.rows[0].cells[2 * i]
        _shade_cell(cell, theme.FILL_CARD)
        for edge_name in ("top", "left", "bottom", "right"):
            _set_cell_border(cell, edge_name, "single", 4, theme.BORDER)

        _para(cell, card["label"], size=theme.SIZE_KPI_LABEL,
              color=theme.TEXT_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=Pt(3), first=True)
        accent = card.get("accent")
        accent_hex = accent.lstrip("#") if accent else None
        _para(cell, card["value"], size=theme.SIZE_KPI_VALUE, bold=True,
              color=accent_hex or theme.TEXT_STRONG,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
        progress = _progress_pct(card)
        if progress is not None:
            fill, achieved = progress
            _card_progress_bar(cell, fill, accent_hex or theme.ACCENT, int(card_w))
            _para(cell, f"목표 대비 {achieved:.1f}% 달성",
                  size=theme.SIZE_CAPTION, color=theme.TEXT_MUTED,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(2))
        if card.get("diff"):
            _para(cell, card["diff"], size=theme.SIZE_KPI_DIFF, bold=True,
                  color=theme.diff_color(card.get("diff_value")) or theme.GRAY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(4))
    _para(document, "", space_after=Pt(4))  # breathing room below the row


# ── data tables ────────────────────────────────────────────────────────

def _is_bar_cell(value):
    return isinstance(value, dict) and value.get("type") == "bar"


def _is_rich_cell(value):
    return isinstance(value, dict) and "text" in value


def _is_image_cell(value):
    return isinstance(value, dict) and value.get("type") == "image"


def _render_image_cell(cell, data_url, width_cm=2.2):
    """Embed a `data:image/...;base64,...` thumbnail (get_ad_creative_info's
    thumbnail_image_data_url) into the cell; falls back to '-' when the
    payload is missing or unreadable."""
    try:
        header, encoded = str(data_url).split(",", 1)
        assert header.startswith("data:image/")
        blob = io.BytesIO(base64.b64decode(encoded))
        paragraph = cell.paragraphs[0]
        paragraph.add_run().add_picture(blob, width=Cm(width_cm))
    except Exception:
        _para(cell, "-", size=theme.SIZE_TD, color=theme.TEXT_FAINT, first=True)


def _cell_text(value):
    if _is_bar_cell(value):
        return str(value.get("label", ""))
    if _is_rich_cell(value):
        return str(value["text"])
    if _is_image_cell(value):
        return ""
    return str(value)


def _display_width(value):
    """Rendered width in half-width character units — CJK glyphs count
    double, so Korean name columns get the room they need."""
    if _is_image_cell(value):
        return 13  # fixed thumbnail width (~2.2cm)
    return sum(2 if ord(ch) > 0x2E80 else 1 for ch in _prettify(_cell_text(value)))


def _column_units(headers, rows):
    """Per-column content width in half-em display units (header included)."""
    units = []
    for c, header in enumerate(headers):
        longest = _display_width(header)
        for row in rows:
            if c < len(row):
                longest = max(longest, _display_width(row[c]))
        units.append(max(longest, 4))
    return units


_PT_PER_UNIT = 0.5      # one half-width glyph ≈ half an em
_CELL_SIDE_PT = 11      # 2 × 110 dxa cell margins per column
_FIT_SAFETY = 1.08


def _fitted_font_pt(units, avail_width):
    """Largest body-font size at which every column renders on one line."""
    avail_pt = int(avail_width) / 12700 - len(units) * _CELL_SIDE_PT
    return avail_pt / (_PT_PER_UNIT * sum(units) * _FIT_SAFETY)


def _column_widths(units, total_width):
    """Distribute the table width proportionally to content units."""
    total_units = sum(units)
    widths = [Emu(int(int(total_width) * u / total_units)) for u in units]
    widths[-1] = Emu(int(total_width) - sum(int(w) for w in widths[:-1]))
    return widths




def _is_total_row(row):
    return bool(row) and _cell_text(row[0]) in _TOTAL_ROW_LABELS


def _is_zero_amount(value):
    text = _cell_text(value).replace("₩", "").replace(",", "").replace("원", "").strip()
    try:
        return float(text) == 0
    except ValueError:
        return False


def _parse_amount(value):
    text = _cell_text(value).replace("₩", "").replace(",", "").replace("원", "").strip()
    try:
        return float(text)
    except ValueError:
        return None


_AD_COST_HEADERS = {"광고비", "광고 비용", "ad_cost"}


def ad_cost_threshold(heading):
    """광고비 필터 기준액 for a section, by heading keyword (키워드 5만 원,
    캠페인/광고그룹 50만 원). None when the section has no threshold."""
    if not heading:
        return None
    for keyword, threshold in theme.AD_COST_FILTERS.items():
        if keyword in str(heading):
            return threshold
    return None


def filter_low_ad_cost(headers, rows, threshold):
    """Drop rows spending below the 광고비 threshold — 합계/총계 rows always
    survive. Returns (rows, removed_count)."""
    cost_col = next((i for i, h in enumerate(headers)
                     if str(h).strip() in _AD_COST_HEADERS), None)
    if cost_col is None or threshold is None:
        return rows, 0
    kept = []
    for row in rows:
        if _is_total_row(row) or cost_col >= len(row):
            kept.append(row)
            continue
        amount = _parse_amount(row[cost_col])
        if amount is None or amount >= threshold:
            kept.append(row)
    return kept, len(rows) - len(kept)


def filter_zero_gross(headers, rows):
    """대용량 표에서 매출(gross) 0원 행 제외 — only kicks in above
    ZERO_GROSS_FILTER_MIN_ROWS so small summary tables stay complete.
    합계/총계 rows are always kept. Returns (rows, removed_count)."""
    if len(rows) <= theme.ZERO_GROSS_FILTER_MIN_ROWS:
        return rows, 0
    gross_col = next((i for i, h in enumerate(headers)
                      if str(h).strip() in _GROSS_HEADERS), None)
    if gross_col is None:
        return rows, 0
    kept = [row for row in rows
            if _is_total_row(row)
            or gross_col >= len(row)
            or not _is_zero_amount(row[gross_col])]
    if not any(not _is_total_row(row) for row in kept):
        # every body row is zero-gross (전환 미집계 브랜드) — an empty table
        # is worse than an unfiltered one, so keep the original rows
        return rows, 0
    return kept, len(rows) - len(kept)


def _truncate_rows(rows, limit):
    """Keep the first rows, always preserving a trailing 합계/총계 row."""
    if len(rows) <= limit:
        return rows, 0
    total_row = None
    body = rows
    if _is_total_row(rows[-1]):
        total_row, body = rows[-1], rows[:-1]
    keep = limit - (1 if total_row is not None else 0)
    shown = body[:keep]
    hidden = len(body) - keep
    if total_row is not None:
        shown = shown + [total_row]
    return shown, hidden


def _set_repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _render_bar_cell(cell, pct, color, label):
    """Approximate a CSS `width:{pct}%` progress bar with a 2-cell nested
    table (filled + track); Word has no native proportional-width bar."""
    pct = max(0.0, min(float(pct), 100.0))
    cell.text = ""
    bar_table = cell.add_table(rows=1, cols=2)
    bar_table.autofit = False
    filled_width = Emu(int(_BAR_WIDTH_EMU * pct / 100))
    track_width = _BAR_WIDTH_EMU - filled_width
    filled_cell, track_cell = bar_table.rows[0].cells
    bar_table.columns[0].width = filled_width
    bar_table.columns[1].width = track_width
    filled_cell.width = filled_width
    track_cell.width = track_width
    _shade_cell(filled_cell, color)
    _shade_cell(track_cell, theme.BORDER)
    _para(cell, label, size=Pt(8), color=theme.TEXT_MUTED, space_before=Pt(2))


def add_data_table(document, heading, headers, rows, rows_total=None,
                   context=None):
    """context: the section heading when the banner is rendered by the
    caller — used for the per-section 광고비 threshold lookup."""
    original_count = rows_total if rows_total is not None \
        else sum(1 for row in rows if not _is_total_row(row))
    threshold = ad_cost_threshold(context or heading)
    has_cost_col = any(str(h).strip() in _AD_COST_HEADERS for h in headers)
    cost_removed = zero_removed = 0
    if threshold and has_cost_col:
        rows, cost_removed = filter_low_ad_cost(headers, rows, threshold)
    else:
        rows, zero_removed = filter_zero_gross(headers, rows)
    shown, _ = _truncate_rows(rows, theme.MAX_TABLE_ROWS)
    body_shown = sum(1 for row in shown if not _is_total_row(row))
    hidden = max(original_count - cost_removed - zero_removed - body_shown, 0)

    # fit check: the landscape page is wide, but a pathological table can
    # still need a slightly smaller size to keep one line per cell
    units = _column_units(headers, shown)
    fitted = _fitted_font_pt(units, theme.CONTENT_W)
    td_size = Pt(round(min(theme.SIZE_TD.pt, max(fitted, theme.MIN_TABLE_FONT)), 1))
    th_size = Pt(max(td_size.pt - 0.5, 7.5))
    content_w = theme.CONTENT_W

    if heading:
        add_heading(document, heading)

    table = document.add_table(rows=1 + len(shown), cols=len(headers))
    _set_table_borders(
        table,
        bottom=("single", 4, theme.BORDER),
        insideH=("single", 4, theme.BORDER_SOFT),
    )
    _set_cell_margins(table)
    _fixed_layout(table, _column_widths(units, content_w))

    header_row = table.rows[0]
    _set_repeat_header_row(header_row)
    for col, header in enumerate(headers):
        cell = header_row.cells[col]
        _shade_cell(cell, theme.FILL_HEADER)
        _set_cell_border(cell, "bottom", "single", 6, theme.BORDER)
        _para(cell, str(header), size=th_size, color=theme.TEXT_TH,
              bold=True, first=True)

    for row_idx, row in enumerate(shown, start=1):
        row_cells = table.rows[row_idx].cells
        is_total = _is_total_row(row)
        for col, value in enumerate(row):
            cell = row_cells[col]
            if _is_bar_cell(value):
                _render_bar_cell(cell, value["pct"], value.get("color", theme.ACCENT),
                                 value.get("label", f"{value['pct']}%"))
            elif _is_image_cell(value):
                _render_image_cell(cell, value.get("data_url"))
            elif _is_rich_cell(value):
                _para(cell, value["text"], size=td_size,
                      color=value.get("color", theme.TEXT_TABLE),
                      bold=value.get("bold", is_total), first=True)
            else:
                _para(cell, value, size=td_size, color=theme.TEXT_TABLE,
                      bold=is_total or None, first=True)
            if is_total:
                _shade_cell(cell, theme.FILL_HEADER)

    notes = []
    if cost_removed:
        notes.append(f"광고비 {threshold:,}원 미만 {cost_removed}행 제외")
    if zero_removed:
        notes.append(f"매출 0원 {zero_removed}행 제외")
    if hidden:
        notes.append(f"상위 {body_shown}행 표시 · 외 {hidden}행 생략")
    if notes:
        _para(document, " · ".join(notes), size=theme.SIZE_CAPTION,
              color=theme.TEXT_FAINT, space_before=Pt(3))
    return table


# ── text & chart sections ──────────────────────────────────────────────

def _is_subheading(line):
    """Short label lines between paragraphs (Overview / 국내분유 / 커피)
    render as bold subheadings — sentence lines end with a stop."""
    text = line.strip()
    return bool(text) and _display_width(text) <= 30 \
        and not text.endswith((".", "!", "?", "%"))


def add_text_section(document, heading, body):
    if heading:
        add_heading(document, heading)
    lines = str(body).split("\n")
    for i, line in enumerate(lines):
        if _is_subheading(line):
            p = _para(document, line, size=theme.SIZE_SUBHEAD, color=theme.TEXT_STRONG,
                      bold=True, space_before=Pt(10) if i else Pt(0), space_after=Pt(2))
        else:
            p = _para(document, line, size=theme.SIZE_BODY, color=theme.TEXT_TABLE,
                      space_after=Pt(5))
            p.paragraph_format.line_spacing = 1.35
        # keep a readable line measure inside the wide landscape page
        p.paragraph_format.left_indent = theme.TEXT_MEASURE_INDENT
        p.paragraph_format.right_indent = theme.TEXT_MEASURE_INDENT


def add_combo_chart_section(document, heading, categories, bar_series, line_series):
    if heading:
        add_heading(document, heading)
    charts.add_combo_chart(document, categories, bar_series, line_series)


def add_line_chart_section(document, heading, categories, series):
    if heading:
        add_heading(document, heading)
    charts.add_line_chart(document, categories, series)
