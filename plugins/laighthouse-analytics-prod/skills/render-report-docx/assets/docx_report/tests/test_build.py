import json
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

import build

_DATA = {
    "title": "다형식품 MTD 보고서",
    "period": "2026-05-01 ~ 2026-05-15",
    "sections": [
        {"type": "heading", "text": "월 목표"},
        {"type": "kpi_cards", "cards": [{"label": "월 목표 매출", "value": "₩ 1,000"}]},
        {"type": "analysis_slot", "key": "section3", "heading": "Executive Summary"},
        {"type": "table", "heading": "목표 달성 현황", "headers": ["지표", "실적"],
         "rows": [["매출", "₩ 500"]]},
        {"type": "chart", "heading": "월별 광고 성과", "categories": ["1월", "2월"],
         "bar_series": [{"name": "광고비", "values": [1, 2]}],
         "line_series": {"name": "ROAS", "values": [3, 4]}},
    ],
}


def test_build_document_all_types_and_heading_merge():
    analysis = {"section3": {"heading": "Executive Summary", "body": "요약 본문."}}
    doc = build.build_document(_DATA, analysis)
    texts = [p.text for p in doc.paragraphs]
    assert "다형식품 MTD 보고서" in texts
    # standalone heading became the kpi heading, with a sequence number
    assert any(t.endswith("월 목표") and t[:2].isdigit() for t in texts)
    assert "요약 본문." in texts        # analysis slot filled
    assert any("목표 달성 현황" in t for t in texts)
    # kpi cards + data table + kpi progress bars don't add doc-level tables
    assert len(doc.tables) == 2


def test_unfilled_analysis_slot_degrades():
    doc = build.build_document(_DATA, analysis=None)
    assert "데이터 준비 중" in [p.text for p in doc.paragraphs]


def test_build_document_rejects_unknown_type():
    with pytest.raises(ValueError, match="unknown section type"):
        build.build_document({"title": "t", "sections": [{"type": "nope"}]})


def test_cli_end_to_end_wellformed(tmp_path):
    data_path = tmp_path / "data.json"
    data_path.write_text(json.dumps(_DATA, ensure_ascii=False), encoding="utf-8")
    analysis_path = tmp_path / "analysis.json"
    analysis_path.write_text(json.dumps(
        {"section3": {"heading": "Executive Summary", "body": "요약."}},
        ensure_ascii=False), encoding="utf-8")
    out_path = tmp_path / "out" / "report.docx"
    subprocess.run(
        [sys.executable, str(Path(build.__file__)),
         "--data", str(data_path), "--analysis", str(analysis_path),
         "--out", str(out_path)],
        check=True,
    )
    assert out_path.exists()
    with zipfile.ZipFile(out_path) as z:
        assert any("charts/chart" in n for n in z.namelist())
        for name in z.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                etree.fromstring(z.read(name))  # raises if malformed
    doc = Document(str(out_path))
    assert any("다형식품" in p.text for p in doc.paragraphs)
