from docx import Document
from lxml import etree

import charts
import sections as sec
import theme

_C = "{http://schemas.openxmlformats.org/drawingml/2006/chart}"
_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

_BARS = [
    {"name": "광고비", "values": [1, 2, 3], "color": "94A3B8"},
    {"name": "전환매출", "values": [4, 5, 6], "color": "93C5FD"},
]
_LINE = {"name": "ROAS(%)", "values": [7, 8, 9], "color": "EF4444"}


def _chart_root():
    d = Document()
    sec.setup_document(d)
    charts.add_combo_chart(d, ["1월", "2월", "3월"], _BARS, _LINE)
    part = next(p for p in d.part.package.iter_parts()
                if "charts/chart" in str(p.partname))
    return d, etree.fromstring(part.blob)


def test_chart_part_created_and_referenced():
    d, root = _chart_root()
    assert root.tag == f"{_C}chartSpace"
    drawings = d.element.body.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr")
    assert len(drawings) == 1


def test_series_child_order_is_schema_valid():
    _, root = _chart_root()
    for ser in root.iter(f"{_C}ser"):
        tags = [etree.QName(ch).localname for ch in ser]
        assert tags.index("tx") < tags.index("spPr")
        assert tags.index("spPr") < tags.index("cat") < tags.index("val")


def test_series_colors_and_axis_numfmt():
    _, root = _chart_root()
    bar_chart = root.find(f".//{_C}barChart")
    fills = [el.get("val") for el in bar_chart.iter(f"{_A}srgbClr")]
    assert "94A3B8" in fills and "93C5FD" in fills
    left_ax = root.findall(f".//{_C}valAx")[0]
    assert left_ax.find(f"{_C}numFmt").get("formatCode") == "#,##0"
    assert left_ax.find(f"{_C}majorGridlines") is not None
    def_rpr = left_ax.find(f"{_C}txPr").find(f".//{_A}defRPr")
    assert def_rpr.find(f"{_A}latin").get("typeface") == theme.FONT


def test_chart_inserted_in_flow_not_after_sectpr():
    d = Document()
    sec.setup_document(d)
    d.add_paragraph("앞 내용")
    charts.add_combo_chart(d, ["a"], [{"name": "b", "values": [1]}],
                           {"name": "l", "values": [2]})
    d.add_paragraph("뒤 내용")
    body = d.element.body
    tags = [etree.QName(ch).localname for ch in body]
    # drawing paragraph must sit before sectPr, between the two text paragraphs
    assert tags[-1] == "sectPr"
    drawing_idx = next(i for i, ch in enumerate(body)
                       if ch.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"))
    after_idx = next(i for i, ch in enumerate(body)
                     if "뒤 내용" in "".join(ch.itertext()))
    assert drawing_idx < after_idx


def test_line_chart_series_colors_from_palette():
    d = Document()
    sec.setup_document(d)
    charts.add_line_chart(d, ["5/1", "5/2"], [
        {"name": "국내분유", "values": [1, 2]},
        {"name": "커피", "values": [3, 4]},
    ])
    part = next(p for p in d.part.package.iter_parts()
                if "charts/chart" in str(p.partname))
    root = etree.fromstring(part.blob)
    assert root.find(f".//{_C}barChart") is None
    sers = root.findall(f".//{_C}lineChart/{_C}ser")
    assert len(sers) == 2
    fills = [el.get("val") for el in root.iter(f"{_A}srgbClr")]
    assert theme.CHART_LINE_COLORS[0] in fills
    assert theme.CHART_LINE_COLORS[1] in fills
    for ser in sers:
        tags = [etree.QName(ch).localname for ch in ser]
        assert tags.index("tx") < tags.index("spPr") < tags.index("cat") < tags.index("val")


def test_legend_bottom_and_distinct_partnames():
    d = Document()
    charts.add_combo_chart(d, ["a"], [{"name": "b", "values": [1]}],
                           {"name": "l", "values": [2]})
    charts.add_combo_chart(d, ["a"], [{"name": "b", "values": [1]}],
                           {"name": "l", "values": [2]})
    names = sorted(str(p.partname) for p in d.part.package.iter_parts()
                   if "charts/chart" in str(p.partname))
    assert names == ["/word/charts/chart1.xml", "/word/charts/chart2.xml"]
    root = etree.fromstring(
        next(p for p in d.part.package.iter_parts()
             if str(p.partname).endswith("chart1.xml")).blob)
    assert root.find(f".//{_C}legend/{_C}legendPos").get("val") == "b"
