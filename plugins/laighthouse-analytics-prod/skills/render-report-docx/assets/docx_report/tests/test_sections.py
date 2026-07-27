from docx import Document
from docx.oxml.ns import qn

import sections as sec
import theme


def _doc():
    d = Document()
    sec.setup_document(d)
    return d


def test_setup_document_a4_and_korean_font():
    d = _doc()
    s = d.sections[0]
    assert abs(s.page_width - theme.PAGE_W) < 1000   # twips round-trip
    assert abs(s.page_height - theme.PAGE_H) < 1000
    normal = d.styles["Normal"]
    assert normal.font.name == theme.FONT
    rfonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    assert rfonts.get(qn("w:eastAsia")) == theme.FONT


def test_title_block_with_accent_rule():
    d = _doc()
    sec.add_title(d, "다형식품 MTD 보고서", "2026-05-01 ~ 2026-05-15")
    texts = [p.text for p in d.paragraphs]
    assert "다형식품 MTD 보고서" in texts
    assert any("2026-05-01" in t for t in texts)


def test_heading_banner_shading_and_accent_bar():
    d = _doc()
    p = sec.add_heading(d, "목표 달성 현황")
    ppr = p.paragraph_format.element.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    assert shd.get(qn("w:fill")) == theme.FILL_BANNER
    left = ppr.find(qn("w:pBdr")).find(qn("w:left"))
    assert left.get(qn("w:color")) == theme.ACCENT
    assert p.paragraph_format.keep_with_next is True


def test_footer_brand_and_page_field():
    d = _doc()
    sec.add_footer(d)
    footer = d.sections[0].footer
    assert "Engineered by Laighthouse AI" in footer.paragraphs[0].text
    assert footer.paragraphs[0]._p.find(qn("w:fldSimple")) is not None


def test_kpi_cards_layout_colors():
    d = _doc()
    sec.add_kpi_cards(d, [
        {"label": "월 예산 목표", "value": "127,636,364"},
        {"label": "기간 누적 ROAS", "value": "559.62%", "accent": "#7c3aed",
         "diff": "목표 425.43%"},
        {"label": "증감", "value": "1", "diff": "▲ +2%", "diff_value": 2},
    ])
    table = d.tables[0]
    assert len(table.columns) == 5  # 3 cards + 2 spacers
    card0 = table.rows[0].cells[0]
    assert card0.paragraphs[0].text == "월 예산 목표"
    assert card0._tc.find(qn("w:tcPr")).find(qn("w:shd")) is not None
    accent_run = table.rows[0].cells[2].paragraphs[1].runs[0]
    assert str(accent_run.font.color.rgb) == "7C3AED"
    diff_run = table.rows[0].cells[4].paragraphs[2].runs[0]
    assert str(diff_run.font.color.rgb) == theme.GREEN


def test_table_header_band_repeat_and_no_vertical_borders():
    d = _doc()
    sec.add_data_table(d, "캠페인별 성과", ["캠페인", "ROAS"],
                       [["브랜드검색", "727%"], ["합계", "573%"]])
    table = d.tables[-1]
    header_row = table.rows[0]
    assert header_row._tr.find(qn("w:trPr")).find(qn("w:tblHeader")) is not None
    shd = header_row.cells[0]._tc.find(qn("w:tcPr")).find(qn("w:shd"))
    assert shd.get(qn("w:fill")) == theme.FILL_HEADER
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders.find(qn("w:insideV")).get(qn("w:val")) == "nil"
    total_run = table.rows[2].cells[0].paragraphs[0].runs[0]
    assert total_run.font.bold is True


def test_table_prettifies_raw_floats():
    d = _doc()
    sec.add_data_table(d, None, ["매체", "소진율"],
                       [["브랜드검색", "55.88918788518661%"]])
    assert d.tables[-1].rows[1].cells[1].paragraphs[0].text == "55.89%"


def test_column_widths_proportional_with_content():
    headers = ["광고그룹", "노출", "광고비"]
    rows = [["012_브랜드_공통_분유_견과류음료_견과류천국", "1,204,331", "808"]]
    units = sec._column_units(headers, rows)
    widths = sec._column_widths(units, theme.CONTENT_W)
    assert widths[0] > widths[1] > 0
    assert abs(sum(int(w) for w in widths) - int(theme.CONTENT_W)) <= 3


def test_document_is_landscape_throughout():
    d = _doc()
    s = d.sections[0]
    assert s.page_width > s.page_height
    # wide 11-column table stays in the same (only) section, one line per cell
    headers = ["캠페인", "네이버 광고 채널명", "매출", "광고비", "ROAS",
               "노출", "클릭", "CTR", "CPC", "구매", "평균단가"]
    rows = [["01_커피믹스_NVSHOP", "PLINK", "152,667,168", "13,820,664", "1105%",
             "557,783", "16,398", "0.4%", "2,074.77", "149", "65,258"]] * 3
    sec.add_data_table(d, "캠페인별 성과", headers, rows)
    assert len(d.sections) == 1
    widths = [c.width for c in d.tables[0].columns]
    # column widths round-trip through twips, so allow half a twip per column
    assert abs(sum(int(w) for w in widths) - int(theme.CONTENT_W)) \
        <= 635 * len(widths)


def test_numbered_heading_banner():
    d = _doc()
    p = sec.add_heading(d, "목표 달성 현황", "03")
    runs = [(r.text, str(r.font.color.rgb)) for r in p.runs]
    assert runs[0] == ("03", theme.ACCENT)
    assert runs[-1][0] == "목표 달성 현황"


def test_kpi_progress_bar_for_percentage_cards():
    d = _doc()
    sec.add_kpi_cards(d, [
        {"label": "기간 예산대비 소진율", "value": "47.95%", "accent": "#3b82f6",
         "diff": "목표 127,636,364 · 소진 61,196,570"},
        {"label": "기간 누적 ROAS", "value": "559.62%", "accent": "#7c3aed",
         "diff": "목표 425.43%"},
        {"label": "월 매출 목표", "value": "₩ 543,000,000"},  # not a percentage
    ])
    cells = d.tables[0].rows[0].cells
    assert len(cells[0].tables) == 1   # progress bar nested table
    assert len(cells[2].tables) == 1   # ROAS card: value vs 목표 → bar
    assert len(cells[4].tables) == 0   # money card: no bar


def test_progress_pct_math():
    # static goal cards (no diff) show no bar even when the value is a %
    assert sec._progress_pct({"value": "425.43%"}) is None
    assert sec._progress_pct(
        {"value": "47.95%", "diff": "목표 127,636,364 · 소진 61,196,570"}) == (47.95, 47.95)
    # ROAS vs target: fill caps at 100, achieved keeps the real ratio
    fill, achieved = sec._progress_pct({"value": "559.62%", "diff": "목표 425.43%"})
    assert fill == 100.0 and round(achieved, 1) == 131.5
    assert sec._progress_pct(
        {"value": "63.07%", "diff": "목표 543,000,000 · 매출 342"}) == (63.07, 63.07)
    assert sec._progress_pct({"value": "₩ 1,000"}) is None


def test_kpi_card_shows_achieved_label():
    d = _doc()
    sec.add_kpi_cards(d, [{"label": "기간 누적 ROAS", "value": "559.62%",
                           "accent": "#7c3aed", "diff": "목표 425.43%"}])
    cell = d.tables[0].rows[0].cells[0]
    texts = [p.text for p in cell.paragraphs]
    assert any("목표 대비 131.5% 달성" in t for t in texts)


def test_text_section_keeps_reading_measure():
    d = _doc()
    sec.add_text_section(d, None, "본문 문장입니다.")
    p = next(x for x in d.paragraphs if x.text == "본문 문장입니다.")
    # indents round-trip through twips — allow sub-millimeter drift
    assert abs(p.paragraph_format.left_indent - theme.TEXT_MEASURE_INDENT) < 1000
    assert abs(p.paragraph_format.right_indent - theme.TEXT_MEASURE_INDENT) < 1000


def test_ad_cost_threshold_by_heading():
    assert sec.ad_cost_threshold("캠페인별 성과") == 500_000
    assert sec.ad_cost_threshold("광고그룹별 성과") == 500_000
    assert sec.ad_cost_threshold("키워드별 성과") == 50_000
    assert sec.ad_cost_threshold("매체별 예산 소진 현황") is None
    assert sec.ad_cost_threshold(None) is None


def test_low_ad_cost_rows_filtered_with_note():
    headers = ["캠페인", "매출", "광고비"]
    rows = [["큰캠페인", "10,000,000", "₩ 8,420,000"],
            ["작은캠페인", "100,000", "₩ 499,999"],
            ["경계캠페인", "900,000", "500,000"],
            ["합계", "11,000,000", "₩ 9,419,999"]]
    d = _doc()
    sec.add_data_table(d, None, headers, rows, context="캠페인별 성과")
    table = d.tables[-1]
    names = [table.rows[r].cells[0].paragraphs[0].text for r in range(1, len(table.rows))]
    assert names == ["큰캠페인", "경계캠페인", "합계"]  # 50만 미만 제외, 합계 보존
    notes = [p.text for p in d.paragraphs if "제외" in p.text]
    assert notes == ["광고비 500,000원 미만 1행 제외"]


def test_keyword_threshold_is_50k():
    headers = ["키워드", "광고비", "매출"]
    rows = [["비싼키워드", "353,722", "8,057,615"], ["싼키워드", "49,999", "0"]]
    d = _doc()
    sec.add_data_table(d, None, headers, rows, context="키워드별 성과")
    assert len(d.tables[-1].rows) == 2  # header + 비싼키워드


def test_section_heading_page_break():
    d = _doc()
    p1 = sec.add_heading(d, "첫 섹션", "01", page_break=False)
    p2 = sec.add_heading(d, "둘째 섹션", "02", page_break=True)
    assert not p1.paragraph_format.page_break_before
    assert p2.paragraph_format.page_break_before is True


def test_zero_gross_filter_on_large_tables():
    headers = ["키워드", "노출", "매출"]
    rows = [[f"kw{i}", "1", "0" if i % 2 else "1,000"] for i in range(30)]
    rows.append(["합계", "30", "15,000"])
    kept, removed = sec.filter_zero_gross(headers, rows)
    assert removed == 15
    assert all(not sec._is_zero_amount(r[2]) or sec._is_total_row(r) for r in kept)
    assert sec._is_total_row(kept[-1])  # 합계 row survives even if checked


def test_zero_gross_filter_skips_small_tables():
    headers = ["지표", "매출"]
    rows = [["a", "0"], ["b", "100"]]
    kept, removed = sec.filter_zero_gross(headers, rows)
    assert removed == 0 and len(kept) == 2


def test_zero_gross_filter_needs_gross_column():
    headers = ["키워드", "노출"]
    rows = [[f"kw{i}", "0"] for i in range(30)]
    kept, removed = sec.filter_zero_gross(headers, rows)
    assert removed == 0 and len(kept) == 30


def test_large_table_notes_mention_filter_and_truncation():
    d = _doc()
    rows = [[f"kw{i}", "0" if i % 2 else "1,000"] for i in range(150)]
    sec.add_data_table(d, "키워드별 성과", ["키워드", "매출"], rows)
    notes = [p.text for p in d.paragraphs if "제외" in p.text or "생략" in p.text]
    assert notes and "매출 0원 75행 제외" in notes[0]
    assert "생략" in notes[0]  # 75 kept > MAX_TABLE_ROWS(50) → also truncated


def test_text_section_subheadings_bold():
    d = _doc()
    sec.add_text_section(d, "카테고리별 성과 분석",
                         "국내분유\n누적 매출 1위이며 성장했습니다.")
    runs = [r for p in d.paragraphs for r in p.runs]
    subhead = next(r for r in runs if r.text == "국내분유")
    assert subhead.font.bold is True
    body = next(r for r in runs if "누적 매출" in r.text)
    assert not body.font.bold


def test_zero_gross_filter_keeps_rows_when_all_zero():
    # 전환 미집계 브랜드(breezm): every row zero — an empty table is worse
    headers = ["소재", "광고비", "매출"]
    rows = [[f"c{i}", "1,000", "0"] for i in range(30)]
    kept, removed = sec.filter_zero_gross(headers, rows)
    assert removed == 0 and len(kept) == 30
